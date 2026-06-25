"""
Generate Parasparam 2026 Abstract in Word (.docx) format.

Template v9.3 guidelines:
  - First Level Heading:  Arial, 14pt, Bold
  - Second Level Heading: Arial, 12pt, Bold
  - Third Level Heading:  Arial, 11pt, Bold, Italic
  - Body text:            Times New Roman, 11pt
  - Abstract text:        Times New Roman, 10pt, Italic
  - Captions:             Times New Roman, 10pt, Centered
  - 3-page limit
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "EzAssist_Parasparam_2026_Abstract.docx")


def set_run_font(run, font_name, size_pt, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force font for East Asian text too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def add_heading1(doc, text):
    """First Level Heading — Arial, 14pt, Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "Arial", 14, bold=True)
    return p


def add_heading2(doc, text):
    """Second Level Heading — Arial, 12pt, Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, "Arial", 12, bold=True)
    return p


def add_heading3(doc, text):
    """Third Level Heading — Arial, 11pt, Bold, Italic"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, "Arial", 11, bold=True, italic=True)
    return p


def add_body(doc, text, space_after=4):
    """Body text — Times New Roman, 11pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", 11)
    return p


def add_abstract_text(doc, text):
    """Abstract text — Times New Roman, 10pt, Italic"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", 10, italic=True)
    return p


def add_bullet(doc, text, level=0):
    """Bulleted body text — Times New Roman, 11pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(f"\u2022  {text}")
    set_run_font(run, "Times New Roman", 11)
    return p


def add_numbered(doc, number, text):
    """Numbered body text — Times New Roman, 11pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"{number}. ")
    set_run_font(run, "Times New Roman", 11, bold=True)
    run2 = p.add_run(text)
    set_run_font(run2, "Times New Roman", 11)
    return p


def add_bold_then_text(doc, bold_part, text_part, indent=0.25):
    """Bold label followed by normal text — Times New Roman, 11pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run1 = p.add_run(bold_part)
    set_run_font(run1, "Times New Roman", 11, bold=True)
    run2 = p.add_run(text_part)
    set_run_font(run2, "Times New Roman", 11)
    return p


def build_document():
    doc = Document()

    # --- Page margins (narrow to fit 3 pages) ---
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ==========================================
    # TITLE
    # ==========================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run("EzAssist: AI-Powered Knowledge Assistant\nfor Faster Support Issue Resolution")
    set_run_font(run, "Arial", 14, bold=True)

    # AUTHORS
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_before = Pt(2)
    author_p.paragraph_format.space_after = Pt(0)
    run = author_p.add_run("Sabyasachi Mallik")
    set_run_font(run, "Times New Roman", 11)

    entity_p = doc.add_paragraph()
    entity_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    entity_p.paragraph_format.space_before = Pt(0)
    entity_p.paragraph_format.space_after = Pt(0)
    run = entity_p.add_run("Hybrid Cloud Customer Services SSD")
    set_run_font(run, "Times New Roman", 11)

    email_p = doc.add_paragraph()
    email_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email_p.paragraph_format.space_before = Pt(0)
    email_p.paragraph_format.space_after = Pt(8)
    run = email_p.add_run("sabyasachi.mallik@hpe.com")
    set_run_font(run, "Times New Roman", 11, italic=True)

    # ==========================================
    # ABSTRACT (Times New Roman, 10pt, Italic)
    # ==========================================
    add_heading1(doc, "Abstract")

    add_abstract_text(doc,
        "EzAssist is an AI-powered Knowledge Assistant that transforms how HPE Ezmeral "
        "support engineers find and apply troubleshooting information. Instead of manually "
        "searching across multiple Knowledge Base (KB) articles, PDF documents, and "
        "Salesforce records, engineers can ask a question in natural language and receive "
        "consolidated, step-by-step troubleshooting guidance from a single interface. "
        "The system uses a Retrieval-Augmented Generation (RAG) architecture combining "
        "a locally-hosted LLaMA 2 7B large language model with semantic search over "
        "PostgreSQL+pgvector to deliver context-aware answers grounded exclusively in "
        "approved support documentation. Engineers can also upload new PDF-based KB "
        "articles and troubleshooting guides, ensuring the knowledge base grows "
        "continuously. A working Proof of Concept is deployed and validated on an HPE "
        "Ezmeral Kubernetes cluster."
    )

    # ==========================================
    # PROBLEM STATEMENT
    # ==========================================
    add_heading1(doc, "Problem Statement")

    add_body(doc,
        "Support Engineers spend a significant amount of time searching across multiple "
        "Knowledge Base (KB) articles, support documents, troubleshooting guides, release "
        "notes, and internal repositories to diagnose and resolve customer issues. "
        "This creates several challenges:"
    )

    add_heading3(doc, "1. Knowledge is Distributed Across Multiple Sources")
    add_body(doc,
        "Critical troubleshooting information is scattered across KB articles, PDF "
        "documents, product manuals, support notes, and internal documentation. "
        "Engineers often need to search multiple systems before finding the required "
        "information."
    )

    add_heading3(doc, "2. Increased Time to Resolution")
    add_body(doc,
        "A considerable portion of troubleshooting time is spent locating relevant "
        "documentation rather than solving the customer\u2019s problem. This increases "
        "Mean Time to Resolution (MTTR) and directly impacts customer satisfaction."
    )

    add_heading3(doc, "3. Knowledge Discovery Becomes Difficult")
    add_body(doc,
        "As the volume of support content grows, identifying the most relevant article "
        "or procedure becomes increasingly challenging, especially for new engineers "
        "or when handling unfamiliar product areas."
    )

    add_heading3(doc, "4. Repeated Effort Across Support Teams")
    add_body(doc,
        "Different engineers frequently spend time searching for the same information "
        "and recreating troubleshooting steps that already exist in documented knowledge "
        "sources, leading to duplicated effort across global teams."
    )

    add_heading3(doc, "5. Limited Ability to Leverage New Knowledge")
    add_body(doc,
        "Newly created KB articles, field advisories, and troubleshooting guides are "
        "difficult to consolidate into a single searchable knowledge repository that "
        "all engineers can benefit from."
    )

    # ==========================================
    # OUR SOLUTION
    # ==========================================
    add_heading1(doc, "Our Solution")

    add_body(doc,
        "We developed EzAssist \u2014 a Hybrid RAG Knowledge Assistant that serves as a "
        "single-stop intelligent support companion for HPE Ezmeral Support teams. "
        "Instead of manually searching through hundreds of KB articles and documents, "
        "engineers simply ask a question in natural language and receive:"
    )

    add_bullet(doc, "Relevant troubleshooting steps extracted from matching KB articles")
    add_bullet(doc, "Associated commands and procedures required for issue resolution")
    add_bullet(doc, "Consolidated information from multiple knowledge sources (PDFs + Salesforce)")
    add_bullet(doc, "Context-aware answers specific to the HPE Ezmeral product line")
    add_bullet(doc, "Continuous knowledge expansion through PDF document upload")

    # --- Technical Architecture ---
    add_heading2(doc, "Technical Architecture")

    add_bold_then_text(doc,
        "LLM Engine \u2014 LLaMA 2 7B Chat (Q4_K_M quantized, ~4 GB GGUF): ",
        "Runs entirely within the Kubernetes pod via llama-cpp-python with a "
        "4,096-token context window, temperature=0.2 for deterministic output, "
        "top_p=0.9, and repeat_penalty=1.2. No external API calls \u2014 all data "
        "stays within the cluster, addressing data sovereignty requirements."
    )

    add_bold_then_text(doc,
        "Embedding Model \u2014 all-MiniLM-L6-v2 (Sentence-Transformers): ",
        "Generates 384-dimensional dense vector embeddings for document chunks "
        "and user queries. Enables semantic similarity matching that goes beyond "
        "keyword search to understand the meaning of support queries."
    )

    add_bold_then_text(doc,
        "Vector Database \u2014 PostgreSQL 15 + pgvector: ",
        "Stores document chunk embeddings and performs L2 distance-based "
        "nearest-neighbor search. Two-pass retrieval: first finds top-k closest "
        "chunks globally, then expands within the same source document to capture "
        "complete procedures. Auto-detects document style (numbered steps vs. "
        "KB-style Issue/Cause/Resolution) to optimize chunk ordering."
    )

    add_bold_then_text(doc,
        "PDF Ingestion Pipeline: ",
        "Multi-stage processing \u2014 text extraction (pdfminer) \u2192 unicode "
        "normalization \u2192 line deduplication \u2192 hyphenation repair \u2192 n-gram "
        "repeat collapsing \u2192 section-aware chunking (preserving headers and "
        "code blocks) \u2192 overlapping window assembly \u2192 sentence-transformer "
        "embedding \u2192 pgvector storage. CLI commands (kubectl, systemctl, helm, "
        "bdconfig) are preserved exactly as documented."
    )

    add_bold_then_text(doc,
        "Salesforce Knowledge Integration: ",
        "Dual-method search using SOSL (full-text) and SOQL (LIKE-based fallback) "
        "against Knowledge Articles with product-line filtering. Rich-text HTML "
        "bodies are converted to clean text using html2text. Supports both "
        "Session ID and OAuth2 authentication."
    )

    add_bold_then_text(doc,
        "Answer Generation: ",
        "Retrieved context chunks are cleaned and assembled into a prompt with "
        "strict grounding rules \u2014 the model uses only commands found exactly "
        "in the retrieved context and refuses to answer when context is "
        "irrelevant, preventing fabricated troubleshooting procedures."
    )

    add_bold_then_text(doc,
        "Deployment: ",
        "Fully containerized with Docker, deployed on Kubernetes using Helm "
        "charts with init containers, Gunicorn WSGI server with --preload, "
        "and an automated deployment script for rapid cluster provisioning."
    )

    # ==========================================
    # EVIDENCE
    # ==========================================
    add_heading1(doc, "Evidence the Solution Works")

    add_body(doc,
        "The system is deployed as a working POC on an HPE Ezmeral Kubernetes cluster "
        "(K8s v1.30.14, Rocky Linux 8.10) and validated with real support scenarios:"
    )

    add_heading3(doc, "Scenario 1 \u2014 Correct Procedure Retrieval")
    add_body(doc,
        "Query: \u201cWhat are the steps for manual restart ERE?\u201d \u2014 The system correctly "
        "retrieved the 7-step restart procedure from the uploaded KB document including "
        "exact systemctl commands (bds-monitoring, bds-worker, bds-controller) and the "
        "required suspendha pre-step."
    )

    add_heading3(doc, "Scenario 2 \u2014 Answer Reliability")
    add_body(doc,
        "When queried about topics not covered in the knowledge base, the system "
        "responds with \u201cI don\u2019t have enough information in the knowledge base to "
        "answer this question\u201d instead of generating plausible-sounding but incorrect "
        "procedures \u2014 a critical safety requirement for production support."
    )

    add_heading3(doc, "Scenario 3 \u2014 Product-Line Filtering")
    add_body(doc,
        "Active filtering ensures only relevant HPE Ezmeral Container Platform articles "
        "are returned, excluding articles from other product families (Data Fabric, "
        "Unified Analytics)."
    )

    # ==========================================
    # COMPETITIVE APPROACHES
    # ==========================================
    add_heading1(doc, "Competitive Approaches")

    # Table: 3 columns
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["Approach", "Advantages", "Disadvantages"]
    rows_data = [
        ["Generic Cloud LLMs\n(ChatGPT, Copilot)",
         "High quality generation;\neasy to use",
         "Data leaves HPE network;\nno SFDC integration;\nhallucination risk"],
        ["HPE GreenLake AI / PCAI",
         "Enterprise-grade\nGPU inference",
         "Requires GPU infrastructure;\nnot targeted at support workflows"],
        ["Manual Salesforce Search",
         "Accurate when engineers\nknow what to look for",
         "Slow; no semantic search;\nno cross-source correlation"],
        ["Generic RAG Frameworks\n(LangChain)",
         "Flexible development\nframework",
         "Single-source only;\nno SFDC integration;\nno product-line filtering"],
    ]

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, "Arial", 10, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows_data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run, "Times New Roman", 9)

    # Caption above table
    # (added after table, but we insert spacing)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(2)
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap.add_run(
        "EzAssist uniquely combines: natural language search across PDFs and live "
        "Salesforce Knowledge Articles, fully offline LLM with no data exfiltration, "
        "answers grounded exclusively in approved documentation, continuous knowledge "
        "expansion through PDF upload, and Helm-packaged deployment for any Kubernetes cluster."
    )
    set_run_font(run, "Times New Roman", 10, italic=True)

    # ==========================================
    # CURRENT STATUS
    # ==========================================
    add_heading1(doc, "Current Status")

    add_body(doc,
        "The solution is a fully functional Proof of Concept deployed on an HPE internal "
        "Kubernetes cluster. It is actively used for HPE Ezmeral Runtime Enterprise "
        "support knowledge retrieval. The Docker image is published (sabya610/rag-app:v2), "
        "the Helm chart is packaged (rag-app-0.2.0.tgz), and the automated deployment "
        "script has been validated for rapid provisioning on new clusters. The web "
        "interface provides three views: Q&A (with PDF/SFDC/Both source selection), "
        "PDF Upload, and Query History."
    )

    # ==========================================
    # NEXT STEPS
    # ==========================================
    add_heading1(doc, "Next Steps")

    add_numbered(doc, 1,
        "GPU-Accelerated Inference: Integrate with HPE Ezmeral PCAI/KServe for "
        "GPU-based LLM inference to reduce response times."
    )
    add_numbered(doc, 2,
        "Automated Salesforce Authentication: Replace manual Session ID management "
        "with OAuth2 Connected App flow for unattended operation."
    )
    add_numbered(doc, 3,
        "User Feedback Loop: Add answer rating (thumbs up/down) to improve retrieval "
        "ranking over time."
    )
    add_numbered(doc, 4,
        "Case History Integration: Ingest case comments and resolution notes from "
        "Salesforce Cases, enabling learning from historical resolutions."
    )
    add_numbered(doc, 5,
        "Multi-Product Support: Parameterize product-line filtering to support other "
        "HPE product families (GreenLake, Aruba, Storage)."
    )

    # ==========================================
    # REFERENCES
    # ==========================================
    add_heading1(doc, "References")

    refs = [
        "[1] P. Lewis et al., \u201cRetrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,\u201d NeurIPS 2020.",
        "[2] Meta AI, \u201cLLaMA 2: Open Foundation and Fine-Tuned Chat Models,\u201d 2023.",
        "[3] llama-cpp-python \u2014 https://github.com/abetlen/llama-cpp-python",
        "[4] pgvector \u2014 https://github.com/pgvector/pgvector",
        "[5] Sentence-Transformers (all-MiniLM-L6-v2) \u2014 https://huggingface.co/sentence-transformers/all-MiniLM-L6-v2",
        "[6] pdfminer.six \u2014 https://github.com/pdfminer/pdfminer.six",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(ref)
        set_run_font(run, "Times New Roman", 9)

    # --- Save ---
    doc.save(OUTPUT_PATH)
    print(f"[OK] Abstract saved to: {OUTPUT_PATH}")
    print(f"     File size: {os.path.getsize(OUTPUT_PATH):,} bytes")
    print("\nNext steps:")
    print("  1. Open in Word, review formatting and page count (must be <= 3 pages)")
    print("  2. Export as PDF (File > Save As > PDF)")
    print("  3. Upload PDF to Parasparam Brightidea portal")


if __name__ == "__main__":
    build_document()
